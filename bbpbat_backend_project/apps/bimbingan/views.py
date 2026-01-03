# File: bbpbat_backend_project/apps/bimbingan/views.py
import os
import io
from docx import Document
from django.conf import settings
from django.utils.dateformat import DateFormat

from rest_framework import generics, permissions, viewsets, status
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework.decorators import action
from django.shortcuts import get_object_or_404
import random
import string
from django.utils import timezone
from django.utils.timesince import timesince
from datetime import timedelta
from django.core.files.base import ContentFile 
from django.db import transaction
from django.db.models import Q



# Import model dan serializer yang dibutuhkan
from .models import Penempatan, Pendaftaran, Laporan, Absensi, Sertifikat
from .serializers import (
    PenempatanSerializer, 
    PendaftaranCreateSerializer, 
    PendaftaranAdminSerializer,
    AbsensiSerializer,
    LaporanAdminSerializer,
    LaporanAdminUpdateSerializer,
    PesertaSertifikatSerializer,
    SertifikatDetailSerializer,
    PenempatanAdminSerializer,  # <-- Tambahkan ini
    PenempatanUpdateSerializer  # 
)

# Import izin kustom dan model dari aplikasi lain
from apps.users.permissions import IsAdminUser
from apps.users.models import User
from apps.peserta.models import PesertaProfile, BuktiPembayaran 

# --- Helper Function yang lebih andal ---
def docx_replace(doc, data):
    """
    Fungsi cerdas untuk mengganti placeholder di seluruh dokumen (.docx),
    termasuk di paragraf, tabel, header, dan footer.
    """
    all_content = doc.paragraphs + \
                  [cell.paragraphs[0] for table in doc.tables for row in table.rows for cell in row.cells]

    for p in all_content:
        for key, value in data.items():
            if key in p.text:
                # Ini adalah trik untuk mengganti teks tanpa merusak format (bold, italic, dll.)
                # dengan memproses setiap 'run' secara terpisah.
                for run in p.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))
    return doc

# --- VIEWS UNTUK PUBLIK (TIDAK DIUBAH) ---
# ... (kode PenempatanListView dan PendaftaranCreateView tetap sama) ...
class PenempatanListView(generics.ListAPIView):
    queryset = Penempatan.objects.all()
    serializer_class = PenempatanSerializer
    permission_classes = [permissions.AllowAny]

class PendaftaranCreateView(generics.CreateAPIView):
    queryset = Pendaftaran.objects.all()
    serializer_class = PendaftaranCreateSerializer
    permission_classes = [permissions.AllowAny]

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)
        headers = self.get_success_headers(serializer.data)
        return Response(
            {"message": "Pendaftaran berhasil diajukan. Silakan tunggu verifikasi dari admin."},
            status=status.HTTP_201_CREATED,
            headers=headers
        )

# --- VIEWSET UNTUK ADMIN ---
# ... (kode AdminPendaftaranViewSet dan AdminDashboardStatsView tetap sama) ...
class AdminPendaftaranViewSet(viewsets.ViewSet):
    permission_classes = [IsAdminUser]

    def list(self, request):
        queryset = Pendaftaran.objects.filter(status=Pendaftaran.Status.PENDING)
        tipe_peserta = request.query_params.get('tipe_peserta', None)
        if tipe_peserta:
            queryset = queryset.filter(tipe_peserta=tipe_peserta)
        serializer = PendaftaranAdminSerializer(queryset, many=True)
        return Response(serializer.data)
    
    @action(detail=True, methods=['post'])
    def approve(self, request, pk=None):
        # ... (implementasi ada)
        pendaftaran = get_object_or_404(Pendaftaran, pk=pk)
        if pendaftaran.status != Pendaftaran.Status.PENDING:
            return Response({"error": "Pendaftaran ini sudah diproses."}, status=status.HTTP_400_BAD_REQUEST)
        password = ''.join(random.choices(string.ascii_letters + string.digits, k=12))
        try:
            user = User.objects.create_user(email=pendaftaran.email, password=password, nama_lengkap=pendaftaran.nama_lengkap, role=User.Role.PESERTA)
        except Exception as e:
            return Response({"error": f"Gagal membuat user: {e}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        PesertaProfile.objects.create(
            user=user, 
            tipe_peserta=pendaftaran.tipe_peserta, 
            nama_institusi=pendaftaran.nama_institusi, 
            no_telepon=pendaftaran.no_telepon, 
            penempatan=pendaftaran.pilihan_penempatan,
            nama_pembimbing=pendaftaran.nama_pembimbing,
            no_telepon_pembimbing=pendaftaran.no_telepon_pembimbing,
            email_pembimbing=pendaftaran.email_pembimbing, 
            )
        pendaftaran.status = Pendaftaran.Status.DISETUJUI
        pendaftaran.user_terkait = user
        pendaftaran.save()
        return Response({"message": "Pendaftaran berhasil disetujui.","email": user.email, "password_sementara": password}, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'])
    def reject(self, request, pk=None):
        # ... (implementasi ada)
        pendaftaran = get_object_or_404(Pendaftaran, pk=pk)
        if pendaftaran.status != Pendaftaran.Status.PENDING:
            return Response({"error": "Pendaftaran ini sudah diproses."}, status=status.HTTP_400_BAD_REQUEST)
        pendaftaran.status = Pendaftaran.Status.DITOLAK
        pendaftaran.save()
        return Response({"message": "Pendaftaran telah ditolak."}, status=status.HTTP_200_OK)

class AdminDashboardStatsView(APIView):
    permission_classes = [IsAdminUser]

    def get(self, request):
        # 1. Ambil data statistik (seperti sebelumnya)
        stats_data = {
            'pendaftar_baru': Pendaftaran.objects.filter(status='PENDING').count(),
            'peserta_aktif': PesertaProfile.objects.filter(status='AKTIF').count(),
            'peserta_selesai': PesertaProfile.objects.filter(status='SELESAI').count(),
            'peserta_lulus': PesertaProfile.objects.filter(status='LULUS').count(),
            'laporan_baru': Laporan.objects.filter(status_review='BARU').count(),
            'laporan_direview': Laporan.objects.filter(status_review='DIREVIEW').count(),
            'laporan_diterima': Laporan.objects.filter(status_review='DITERIMA').count(),
        }

        # 2. Ambil 5 pendaftaran terbaru yang masih pending
        pendaftaran_terbaru = Pendaftaran.objects.filter(status='PENDING').order_by('-tanggal_daftar')[:5]
        
        # 3. Ambil 5 laporan terbaru yang statusnya 'BARU'
        laporan_terbaru = Laporan.objects.select_related('profil__user').filter(status_review='BARU').order_by('-disubmit_pada')[:5]

        # 4. Gabungkan dan format aktivitas terbaru
        aktivitas = []
        for p in pendaftaran_terbaru:
            aktivitas.append({
                'tipe': 'PENDAFTARAN',
                'teks': f"Pendaftaran baru dari {p.nama_lengkap}",
                'waktu_raw': p.tanggal_daftar,
                'waktu': f"{timesince(p.tanggal_daftar)} yang lalu"
            })
        
        for l in laporan_terbaru:
            aktivitas.append({
                'tipe': 'LAPORAN',
                'teks': f"Laporan baru dari {l.profil.user.nama_lengkap}",
                'waktu_raw': l.disubmit_pada,
                'waktu': f"{timesince(l.disubmit_pada)} yang lalu"
            })
        
        # 5. Urutkan semua aktivitas berdasarkan waktu, ambil 5 teratas
        aktivitas_terurut = sorted(aktivitas, key=lambda x: x['waktu_raw'], reverse=True)[:5]

        # 6. Kirim semua data dalam satu respons
        response_data = {
            'stats': stats_data,
            'aktivitas_terbaru': aktivitas_terurut
        }
        
        return Response(response_data)

# =================================================================
#             VIEWSET UNTUK MANAJEMEN ABSENSI ADMIN
# =================================================================
class AbsensiAdminViewSet(viewsets.ModelViewSet):
    """
    ViewSet untuk admin mengelola data absensi peserta.
    - Mendukung filter berdasarkan tanggal (?tanggal=YYYY-MM-DD).
    - Mendukung filter berdasarkan tipe peserta (?tipe_peserta=student atau ?tipe_peserta=general).
    """
    serializer_class = AbsensiSerializer
    permission_classes = [IsAdminUser]

    def get_queryset(self):
        """
        Override method ini untuk optimasi query dan implementasi filter.
        """
        # 1. Query dasar dengan optimasi (menghindari N+1 problem)
        # Mengambil semua data terkait dalam satu query database.
        queryset = Absensi.objects.select_related(
            'profil__user', 
            'profil__penempatan'
        ).all()

        # 2. Ambil parameter filter dari URL
        tanggal_filter = self.request.query_params.get('tanggal', None)
        tipe_peserta_filter = self.request.query_params.get('tipe_peserta', None)

        # 3. Terapkan filter jika ada
        if tanggal_filter:
            queryset = queryset.filter(tanggal=tanggal_filter)
        
        if tipe_peserta_filter:
            # Konversi 'student'/'general' dari frontend ke format model
            tipe_model = 'PELAJAR' if tipe_peserta_filter.lower() == 'student' else 'UMUM'
            queryset = queryset.filter(profil__tipe_peserta=tipe_model)

        # Urutkan berdasarkan nama untuk tampilan yang konsisten
        return queryset.order_by('profil__user__nama_lengkap')


# =================================================================
#                 VIEWSET UNTUK MANAJEMEN LAPORAN ADMIN
# =================================================================
class LaporanAdminViewSet(viewsets.ModelViewSet):
    """
    ViewSet untuk admin mengelola laporan peserta.
    - Mendukung filter berdasarkan tipe peserta dan status review.
    - Menggunakan serializer yang berbeda untuk membaca (list/detail) dan menulis (update).
    """
    permission_classes = [IsAdminUser]

    def get_queryset(self):
        """
        Query dasar dengan optimasi dan filter.
        """
        queryset = Laporan.objects.select_related(
            'profil__user', 
            'profil__penempatan'
        ).all()

        # Filter berdasarkan tipe peserta (student/general)
        tipe_peserta_filter = self.request.query_params.get('tipe_peserta', None)
        if tipe_peserta_filter:
            # tipe_model = 'PELAJAR' if tipe_peserta_filter.lower() == 'student' else 'UMUM'
            queryset = queryset.filter(profil__tipe_peserta=tipe_peserta_filter)

        # Filter berdasarkan status review (baru, direview, dll.)
        status_filter = self.request.query_params.get('status', None)
        if status_filter and status_filter != 'all':
            # Frontend mengirim huruf kecil, model menggunakan huruf besar
            queryset = queryset.filter(status_review=status_filter.upper())

        return queryset.order_by('-disubmit_pada')

    def get_serializer_class(self):
        """
        Gunakan serializer yang berbeda untuk aksi yang berbeda.
        - LaporanAdminSerializer (kaya data) untuk 'list' dan 'retrieve'.
        - LaporanAdminUpdateSerializer (ramping) untuk 'update' dan 'partial_update'.
        """
        if self.action in ['update', 'partial_update']:
            return LaporanAdminUpdateSerializer
        return LaporanAdminSerializer

class AdminSertifikatViewSet(viewsets.ViewSet):
    """
    ViewSet final untuk manajemen sertifikat oleh Admin.
    - list: Menampilkan daftar peserta yang layak.
    - create: Menerbitkan sertifikat secara aman menggunakan template Word.
    """
    permission_classes = [IsAdminUser]

    def list(self, request):
        """
        Menampilkan daftar peserta yang memenuhi syarat sertifikat
        berdasarkan aturan baru.
        """
        # Syarat untuk peserta PELAJAR
        syarat_pelajar = Q(
            tipe_peserta=PesertaProfile.TipePeserta.PELAJAR,
            laporan__status_review=Laporan.StatusReview.DITERIMA
        )
        
        # Syarat untuk peserta UMUM
        # Asumsi status pembayaran yang valid adalah "Telah Diverifikasi"
        syarat_umum = Q(
            tipe_peserta=PesertaProfile.TipePeserta.UMUM,
            pembayaran__status_verifikasi="Telah Diverifikasi"
        )
        
        # Query utama yang menggabungkan kedua syarat
        queryset = PesertaProfile.objects.filter(
            syarat_pelajar | syarat_umum,  # <-- Logika ATAU (OR)
            status=PesertaProfile.StatusPeserta.SELESAI,
            sertifikat__isnull=True
        ).distinct().select_related('user', 'penempatan')

        serializer = PesertaSertifikatSerializer(queryset, many=True)
        return Response(serializer.data)
    
    @transaction.atomic
    def create(self, request):
        """ Logika untuk membuat sertifikat menggunakan template PDF. """
        peserta_id = request.data.get('peserta_id')
        if not peserta_id:
            return Response({"detail": "ID Peserta wajib diisi."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            peserta = PesertaProfile.objects.select_for_update().get(id=peserta_id)
        except PesertaProfile.DoesNotExist:
            return Response({"detail": "Peserta tidak ditemukan."}, status=status.HTTP_404_NOT_FOUND)

        # --- VALIDASI PRASYARAT ---
        if hasattr(peserta, 'sertifikat'):
            return Response({"detail": "Peserta ini sudah memiliki sertifikat."}, status=status.HTTP_400_BAD_REQUEST)
        
        if peserta.status != PesertaProfile.StatusPeserta.SELESAI:
            return Response({"detail": "Peserta belum menyelesaikan program."}, status=status.HTTP_400_BAD_REQUEST)

        # =================================================================
        # INI BLOK LOGIKA YANG DIPERBAIKI (PDF VERSION)
        # =================================================================
        if peserta.tipe_peserta == PesertaProfile.TipePeserta.PELAJAR:
            # Syarat untuk Pelajar: Laporan Diterima
            if not Laporan.objects.filter(profil=peserta, status_review=Laporan.StatusReview.DITERIMA).exists():
                return Response({"detail": "Error: Laporan akhir peserta belum diterima."}, status=status.HTTP_400_BAD_REQUEST)
        
        elif peserta.tipe_peserta == PesertaProfile.TipePeserta.UMUM:
            # Syarat untuk Umum: Pembayaran Diverifikasi
            if not BuktiPembayaran.objects.filter(profil=peserta, status_verifikasi="Telah Diverifikasi").exists():
                return Response({"detail": "Error: Bukti pembayaran peserta belum diverifikasi."}, status=status.HTTP_400_BAD_REQUEST)

        # 2. Generate Nomor Sertifikat
        tahun = timezone.now().year
        sertifikat_terakhir = Sertifikat.objects.filter(nomor_sertifikat__startswith=f"BBPBAT/CERT/{tahun}").order_by('id').last()
        nomor_urut = 1
        if sertifikat_terakhir:
            try:
                nomor_urut = int(sertifikat_terakhir.nomor_sertifikat.split('/')[-1]) + 1
            except ValueError:
                pass # Fallback to 1 if parsing fails
        nomor_sertifikat = f"BBPBAT/CERT/{tahun}/{str(nomor_urut).zfill(3)}"

        # 3. Tentukan Path Template PDF
        # Saat ini kita gunakan template yang sama dulu untuk demo, atau sesuaikan nama filenya
        # template_name = 'E_SERTIFIKAT_PKL_MAGANG.pdf' 
        # Atau jika ada beda:
        template_name = 'E_SERTIFIKAT_PKL_MAGANG.pdf' # Gunakan file yang diminta user
        
        template_path = os.path.join(settings.BASE_DIR, 'certificate_templates', template_name)
        if not os.path.exists(template_path):
            return Response({"detail": f"File template PDF tidak ditemukan: {template_name} di {template_path}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # 4. Siapkan Data Pengganti (Konteks)
        data_konteks = {
            "{{NAMA_LENGKAP}}": peserta.user.nama_lengkap or '',
            # "{{NOMOR_INDUK}}": peserta.nomor_induk or 'N/A', # PDF Template mungkin tidak butuh ini, cek pdf_utils
            "{{NAMA_INSTITUSI}}": peserta.nama_institusi or '',
            "{{NOMOR_SERTIFIKAT}}": nomor_sertifikat,
            "{{TANGGAL_TERBIT}}": DateFormat(timezone.now()).format('j F Y'),
            "{{TANGGAL_MULAI}}": DateFormat(peserta.tanggal_mulai).format('j F Y') if peserta.tanggal_mulai else 'N/A',
            "{{TANGGAL_SELESAI}}": DateFormat(peserta.tanggal_selesai).format('j F Y') if peserta.tanggal_selesai else 'N/A',
            "{{PENEMPATAN}}": peserta.penempatan.nama if peserta.penempatan else 'N/A',
        }
        
        # 5. Import dan Panggil Generator PDF
        # Pastikan import ada di atas file, atau import di sini (lazy import) untuk menghindari circular imports di views besar
        from .pdf_utils import generate_pdf_certificate
        
        try:
            pdf_stream = generate_pdf_certificate(template_path, data_konteks)
        except Exception as e:
            return Response({"detail": f"Gagal generate PDF: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # 6. Buat dan Simpan Objek Sertifikat
        sertifikat = Sertifikat.objects.create(
            profil=peserta,
            nomor_sertifikat=nomor_sertifikat,
            template_digunakan=Sertifikat.TemplateType.PELAJAR if peserta.tipe_peserta == PesertaProfile.TipePeserta.PELAJAR else Sertifikat.TemplateType.UMUM
        )
        file_name = f'Sertifikat_{peserta.user.nama_lengkap.replace(" ", "_")}.pdf'
        sertifikat.file_sertifikat.save(file_name, pdf_stream, save=True)
        
        # 7. Update status peserta
        peserta.status = PesertaProfile.StatusPeserta.LULUS
        peserta.save(update_fields=['status'])
        
        serializer = SertifikatDetailSerializer(sertifikat, context={'request': request})
        return Response(serializer.data, status=status.HTTP_201_CREATED)

# =================================================================
# ===== TAMBAHKAN VIEWSET BARU DI BAGIAN AKHIR UNTUK KUOTA ========
# =================================================================
class PenempatanAdminViewSet(viewsets.ModelViewSet):
    """
    ViewSet untuk Admin mengelola unit penempatan dan kuotanya.
    - list (GET): Menampilkan semua unit dengan detail kuota terisi.
    - partial_update (PATCH): Memperbarui kuota dan mengembalikan data lengkap.
    """
    permission_classes = [IsAdminUser]
    queryset = Penempatan.objects.all().order_by('nama')

    def get_serializer_class(self):
        """
        Menggunakan serializer yang berbeda untuk validasi dan response.
        """
        if self.action in ['update', 'partial_update']:
            # Gunakan serializer ramping ini HANYA untuk validasi data yang masuk
            return PenempatanUpdateSerializer
        
        # Gunakan serializer kaya data ini untuk semua response ke frontend
        return PenempatanAdminSerializer

    def update(self, request, *args, **kwargs):
        """
        Override method 'update' (PUT & PATCH) untuk memastikan response
        selalu mengembalikan data yang lengkap dan ter-serialisasi dengan benar.
        """
        # Jalankan proses update standar
        partial = kwargs.pop('partial', False)
        instance = self.get_object()
        # Serializer validasi (UpdateSerializer) digunakan di sini secara internal
        serializer = self.get_serializer(instance, data=request.data, partial=partial)
        serializer.is_valid(raise_exception=True)
        self.perform_update(serializer)

        # Setelah update berhasil, buat serializer BARU dengan data yang kaya (AdminSerializer)
        # untuk dikirim sebagai response.
        response_serializer = PenempatanAdminSerializer(instance, context={'request': request})
        return Response(response_serializer.data)